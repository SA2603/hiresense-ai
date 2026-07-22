import pdfplumber
import docx
import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts all readable text from a PDF file.
    file_bytes: the raw bytes of the uploaded PDF.
    Returns a single string with all pages' text joined together.
    """
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:  # some pages might have no extractable text (e.g. scanned images)
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts all paragraph text from a DOCX file.
    file_bytes: the raw bytes of the uploaded DOCX.
    """
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]

    # Also extract text inside tables - many resumes use tables for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text(uploaded_file) -> str:
    """
    Detects file type from a Streamlit UploadedFile object and routes
    to the correct extractor. Returns the extracted plain text, or
    raises a ValueError for unsupported file types.
    """
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF or DOCX.")

    if not text.strip():
        raise ValueError(
            "No text could be extracted. The file may be a scanned image "
            "or empty — please upload a text-based PDF or DOCX."
        )

    return text.strip()